import streamlit as st
import pandas as pd
import io
from docx import Document

def process_board_decisions(excel_file, programme_name, sheet_name="Summary"):
    """
    Processes board decisions from an Excel file-like object and creates a Word document.
    
    Expected column layout on the 'Summary' sheet:
      - Column C: Student ID (index 2)
      - Column I: Decision/Notes (index 8)
      - Column E: Name (index 4)
      - Column L: Template Code (index 11)
    
    The programme name is provided manually.
    Returns a tuple: (doc_bytes, list_of_decision_texts)
    """
    try:
        df = pd.read_excel(excel_file, sheet_name=sheet_name)
    except Exception as e:
        st.error(f"Error reading the sheet '{sheet_name}': {e}")
        return None, None

    doc = Document()
    decision_texts = []  # To collect generated decision texts

    for index, row in df.iterrows():
        studentID = row.iloc[2]   # Column C: Student ID
        decision  = row.iloc[8]    # Column I: Decision/Notes
        name      = row.iloc[4]    # Column E: Name
        templateCode = row.iloc[11] if len(row) > 11 else ""
        if pd.isna(templateCode):
            templateCode = ""

        if templateCode == "" or templateCode is None:
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "There is no need to paste this record into the SPI screen in SITS."
        elif templateCode == "A":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              f"{decision}\n\n" \
                              "The deadline for your reassessment(s) and further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        elif templateCode == "B":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              "Due to the volume of reassessments, the Board has recommended that you return next academic year. " \
                              "You will need to meet with your academic tutor to discuss your adjusted deadlines and workload.\n\n" \
                              f"{decision}\n\n" \
                              "Further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        elif templateCode == "C":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              "Due to the volume of reassessments, the Board has recommended that you complete some of your reassessments in August " \
                              "and return next academic year to complete the rest.\n\n" \
                              f"{decision}\n\n" \
                              "The deadline for your August reassessment(s) and further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You will need to meet with your academic tutor to discuss your adjusted deadlines and workload. " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        elif templateCode == "D":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile. " \
                              "You have reached the maximum capacity for reassessment, and so cannot continue on the programme.\n\n" \
                              "Further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you understand the decision that has been made and the support that is available."
        elif templateCode == "E":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              "You are no longer eligible for a Bachelor’s degree. However, the Board has offered you an optional reassessment opportunity " \
                              "in order to be eligible for a Certificate/Diploma of Higher Education:\n\n" \
                              f"{decision}\n\n" \
                              "The deadline for your reassessment(s) and further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        elif templateCode == "F":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              "You have achieved a provisional award of [AWARD]. However, you are being offered re-assessments which may improve your degree classification/grade profile.\n\n" \
                              "[DECISION]\n\n" \
                              "The deadline for your reassessment(s) and further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        elif templateCode == "G":
            templateContent = f"{studentID}, {name}, {programme_name}\n\n" \
                              "Your programme’s Board of Examiners recently met to review your grade profile and the outcome is as follows.\n\n" \
                              "You are unable to achieve a Bachelor’s degree with honours. To achieve an Ordinary Bachelor’s degree, the Board of Examiners " \
                              "requires you to take the following compulsory re-assessment(s):\n\n" \
                              f"{decision}\n\n" \
                              "The deadline for your reassessment(s) and further guidance is provided on the College’s Results Service webpage - " \
                              "https://students.brunel.ac.uk/study/chmls/results-service (Please copy and paste this link into your browser). " \
                              "You must check this page so you gain an understanding of your current grades and the necessary steps of what to do next."
        else:
            templateContent = "No valid template code provided."

        doc.add_paragraph(templateContent)
        doc.add_paragraph("-" * 40)
        decision_texts.append(templateContent)
    
    # Save document to a BytesIO stream
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_bytes = doc_io.getvalue()
    
    return doc_bytes, decision_texts

# -------------------- Streamlit App -------------------- #

st.title("Progression Note Generator")

st.markdown("""
This app generates progression notes for students based on the Exam Board Spreadsheet.
**User Instructions:**
- **Template codes** should be entered into **Column L** on the Exam Board Spreadsheet.
- Manually enter the **Programme Name** below.
- Upload the **Exam Board Spreadsheet** (Excel file with `.xlsx` extension).
- A Word document will be generated with the progression notes and will be available for download.
- The generated decisions will also be displayed on the page.
""")

# User input for Programme Name
programme_name = st.text_input("Enter Programme Name:")

# File uploader for the Exam Board Spreadsheet
uploaded_file = st.file_uploader("Upload Exam Board Spreadsheet", type="xlsx")

if uploaded_file and programme_name:
    doc_bytes, decisions = process_board_decisions(uploaded_file, programme_name, sheet_name="Summary")
    
    if doc_bytes is not None:
        # Generate file name based on the programme name
        doc_filename = f"{programme_name}_Board_Decisions.docx"
        st.success("Document generated successfully!")
        
        # Download button for the generated Word document
        st.download_button(
            label="Download Generated Word Document",
            data=doc_bytes,
            file_name=doc_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Display the generated decisions on the page
        st.header("Generated Decisions:")
        for decision_text in decisions:
            st.markdown(decision_text)
            st.markdown("---")
else:
    st.info("Please enter the Programme Name and upload the Exam Board Spreadsheet to generate the progression notes.")

st.header("Template Codes Reference")
template_data = {
    "Template Code": ["(Empty)", "A", "B", "C", "D", "E", "F", "G"],
    "Description": [
        "No template code provided. Record does not need to be pasted into SPI screen in SITS.",
        "Board review with grade profile outcome. Inserts decision. Check Results Service webpage for guidance.",
        "Board review with recommendation to return next academic year. Inserts decision. Check Results Service webpage for guidance.",
        "Board review with recommendation to complete some reassessments in August and return next academic year. Inserts decision. Check Results Service webpage for guidance.",
        "Board review where maximum capacity for reassessment has been reached. Cannot continue on the programme.",
        "Board review where student is no longer eligible for a Bachelor’s degree, but is offered an optional reassessment opportunity for Certificate/Diploma.",
        "Board review with provisional award and opportunity for re-assessments to improve degree classification. Uses [AWARD] and [DECISION] placeholders.",
        "Board review where student cannot achieve a Bachelor’s degree with honours and must take compulsory re-assessment(s) for an Ordinary Bachelor’s degree."
    ]
}
template_df = pd.DataFrame(template_data)
st.table(template_df)

